from unittest import skip

from fastapi import FastAPI, UploadFile, File, HTTPException, Request
from fastapi.responses import FileResponse
from pypdf import PdfReader
from openai import OpenAI
import os
import json
from datetime import date, datetime, timedelta
from docx import Document
from dotenv import load_dotenv

load_dotenv()
print("KEY LOADED:", os.getenv("OPENAI_API_KEY")[:8])
app = FastAPI()

@app.get("/")
async def read_index():
    return FileResponse('templates/index.html') 

# --- COPY THESE FUNCTIONS FROM app.py AS-IS ---
def generate_ai_today_text(syllabus_text, client):
    prompt = f"""
Read the syllabus text below and extract all important academic deadlines and events that should be added to a student's personal calendar.

Look for:
- assignments
- quizzes
- exams
- projects
- presentations
- labs
- essays
- reports
- readings only if they have a due date
- important class deadlines
- final exams or major assessments
- labs only when something must be submitted, completed, or prepared by a certain date
- waivers or required forms only when they have a due date
- graded attendance only if the syllabus clearly treats it as an assessed requirement with a date

Do NOT include in Deadlines Found:
- lecture topics
- attendance
- weekly themes
- class meetings
- field trips
- artist talks
- readings unless explicitly due
- general course activities
- office hours
- location information
- announcements without a deadline

Return only valid JSON.
Do not return explanations, markdown, headings, or extra text.

Return a JSON array.
Each item in the array must follow this format:
{{
  "title": "string",
  "type": "assignment | quiz | exam | project | presentation | lab | reading | reflection | paper | essay | report | lab | deadline",
  "date": "YYYY-MM-DD or null",
  "time": "HH:MM or null",
  "description": "short helpful detail or null",
  "source_section": "evaluation | schedule | other",
  "evidence_text": "exact short quote or excerpt from the syllabus supporting the date"
}}

Rules:
- Prioritize course-related deadlines over administrative/university-wide deadlines
- Only include items that have a clear due date or scheduled date.
- Do not guess missing dates or times.
- If the date is missing, do not include the item.
- If the time is missing, use null.
- Keep titles short and student-friendly.
- Sort items by date from earliest to latest.
- Don't include items that are not calendar-worthy (e.g., "Office Hours", "Syllabus Overview")
- If a syllabus contains both a course evaluation section and a weekly schedule, prioritize dates from the evaluation/grading section for quizzes, exams, projects, presentations, reflections, and other graded work.
- Do not use lecture-topic dates as assignment dates unless the syllabus clearly says the item is due on that date.
- For each item, label where the date came from:
  - "evaluation" if it came from a grading/evaluation/assessment table or section
  - "schedule" if it came from the weekly course calendar or class schedule
  - "other" otherwise
- If the same graded item appears more than once with conflicting dates, prefer the one from "evaluation".
- Only include an item if you can provide a short exact supporting excerpt from the syllabus.
- If the date is uncertain or conflicting, prefer the item with clearer evidence from an evaluation/grading section.
- If no valid calendar items are found, return [].
- Classification rules for "type":
  - Only classify an item as "exam" if it explicitly includes words like "exam", "midterm", "test", or "quiz"
  - Do NOT classify an item as "exam" just because it contains the word "final"
  - Items such as "final project", "final paper", "final portfolio", "final assignment" must be classified as "assignment"
  - If an item contains both "final" and assignment-related words (e.g., project, paper, portfolio, assignment), classify it as "assignment"
  - Words like "assignment", "project", "paper", "essay", "portfolio", "submission", "reflection", "plan", "create" should strongly indicate type "assignment"

Syllabus text:
{syllabus_text}
"""
    response = client.responses.create(
        model="gpt-4.1-mini",
        input=prompt
    )
    return response.output_text

def generate_prep_events(title, due_date, task_type):
    prep_events = [] # Create an empty list where I will store all the prep events
    clean_title = title.replace("Due", "").strip()
    clean_title = clean_title.replace("Research Draft", "Proposal")
    clean_title = clean_title.replace("Take-Home Exam", "Exam")
    # Step 1: decide prep steps 
    if task_type == "assignment":
        steps = [("Start Draft", 5), ("Final Review", 1)]
    elif task_type == "exam":
        steps = [("Study Session", 5), ("Review", 1)]
    elif task_type == "quiz":
        steps = [("Review Notes", 2)]
    elif task_type == "project":
        steps = [("Start Project", 7), ("Work Session", 3), ("Final Review", 1)]
    elif task_type == "presentation":
        steps = [("Prepare Slides", 4), ("Practice", 1)]
    elif task_type == "reflection":
        steps = [("Review Notes", 2), ("Draft", 1)]
    elif task_type == "lab":
        steps = [("Prepare", 1)]
    elif task_type == "deadline":
        steps = [("Prepare", 3), ("Submit", 1)]
    else:
        steps = []

    due = datetime.strptime(due_date, "%Y-%m-%d")

    for step_name, days_before in steps:
        prep_date = due - timedelta(days=days_before)

        event_title = f"{clean_title} — {step_name}"

        words = event_title.split()
        event_title = " ".join(dict.fromkeys(words))

        prep_events.append({
            "title": event_title,
            "date": prep_date.strftime("%Y-%m-%d")
        }) 
    return prep_events

def prefer_evaluation_dates(calendar_items):
    final = {}
    for item in calendar_items:
        key = item.get("title", "").strip().lower()
        if not key:
            continue

        # If same item appears twice → prefer evaluation
        if key not in final or item.get("source_section") == "evaluation":
            final[key] = item

    return list(final.values())

def group_events(events):
    assignments = []
    exams = []
    others = []

    for event in events:
        event_type = event.get("type", "").lower()
        title = event.get("title", "").lower()

        if any(word in title for word in ["assignment", "project", "essay", "paper", "reflection", "presentation"]):
            assignments.append(event)

        elif any(word in title for word in ["exam", "quiz", "midterm", "final"]):
            exams.append(event)

        elif event_type in ["assignment", "project"]:
            assignments.append(event)

        elif event_type in ["exam", "quiz"]:
            exams.append(event)

        else:
            others.append(event)

    return assignments, exams, others

def sort_events(events):
    return sorted(events, key=lambda x: x["date"])

def validate_output(calendar_items):
    valid_items = []
    for item in calendar_items: #loop through every item 
        errors = [] #creates a fresh error list for each item 
        if not item.get("date"):
            errors.append("Missing Date")
        if not item.get("title"):
            errors.append("Missing Title")
        if not item.get("type"):
            errors.append("Missing Type")
        
        item["errors"] = errors  #attaches the error list
        item["is_valid"] = len(errors) == 0   # quick true/false flag      
        valid_items.append(item) #ship it 

    return valid_items

@app.post("/api/analyze")
async def analyze_syllabus(file: UploadFile = File(...)):
    if not file.filename.endswith((".pdf", ".docx")):
        raise HTTPException(400, "Only PDF & DOCX files accepted")

    syllabus_text = ""

    if file.filename.endswith(".pdf"):
        pdf_reader = PdfReader(file.file)
        for page in pdf_reader.pages:
            page_text = page.extract_text()
            if page_text:
                syllabus_text += page_text + "\n"

    elif file.filename.endswith(".docx"):
        doc = Document(file.file)
        for para in doc.paragraphs:
            if para.text.strip():
                syllabus_text += para.text.strip() + "\n"
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    syllabus_text += " | ".join(row_text) + "\n"

    if not syllabus_text.strip():
        raise HTTPException(400, "Could not extract text from this file")

    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key:
        raise HTTPException(500, "OpenAI API key not found")

    client = OpenAI(api_key=api_key)
    ai_response = generate_ai_today_text(syllabus_text, client)

    try:
        calendar_items = json.loads(ai_response)
        calendar_items = validate_output(calendar_items)
        calendar_items = prefer_evaluation_dates(calendar_items)
        
        unique_events = []
        seen = set()

        for event in calendar_items:
            key = (event["title"], event["date"])
            if key not in seen:
                seen.add(key)
                unique_events.append(event)

        calendar_items = unique_events

        assignments, exams, others = group_events(calendar_items)
        assignments = sort_events(assignments)
        exams = sort_events(exams)
        others = sort_events(others)

    except json.JSONDecodeError:
        raise HTTPException(500, "AI response was not valid JSON")
            
    prep_events_lists = []
    for task in calendar_items:
        title = task["title"]
        due_date = task["date"]
        task_type = task["type"]
        if not due_date:
            continue
        prep_events = generate_prep_events(title, due_date, task_type)
        prep_events_lists.extend(prep_events)

    # Build ICS file text
    calendar_text = "BEGIN:VCALENDAR\nVERSION:2.0\nPRODID:-//Syllabus Parser//EN\n"
    for item in calendar_items:
        title = item["title"]
        date = item["date"]
        if not date:
            continue
        description = item.get("description", "")
        event_text = f"""BEGIN:VEVENT
SUMMARY:{title}
DTSTART;VALUE=DATE:{date.replace("-", "")}
DESCRIPTION:{description if description else ""}
END:VEVENT
"""
        calendar_text += event_text

    for prep_event in prep_events_lists:
        title = prep_event["title"]
        date = prep_event["date"]
        event_text = f"""BEGIN:VEVENT
SUMMARY:{title}
DTSTART;VALUE=DATE:{date.replace("-", "")}
END:VEVENT
"""
        calendar_text += event_text

    calendar_text += "END:VCALENDAR\n"

    return {
        "deadlines": calendar_items,   
        "assignments": assignments,
        "exams": exams,
        "others": others,
        "prep_events": prep_events_lists,
        "ics_content": calendar_text
    }